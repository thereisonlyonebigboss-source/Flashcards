"""
text_extraction.py

Utilities to scan folders and extract text from various file formats.
Supports .txt, .md, .docx, and .pdf files with robust error handling.
"""

from pathlib import Path
from typing import Generator, Dict

import pdfplumber
import docx

# Supported file extensions for v1
SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


def iter_note_files(root_folder: Path) -> Generator[Dict, None, None]:
    """
    Recursively scan root_folder and yield metadata for supported files.

    Args:
        root_folder: Path to the root notes folder

    Yields:
        Dict containing:
        - full_path: Path object to the file
        - subject: Subject name extracted from folder structure
        - subtopic: Subtopic name extracted from folder structure
        - filename: Name of the file
    """
    root_folder = root_folder.resolve()

    for path in root_folder.rglob("*"):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                rel = path.relative_to(root_folder)
                parts = rel.parts

                # Extract subject and subtopic from folder structure
                # Expected: root/Subject/Subtopic/files or root/Subject/files
                if len(parts) >= 2:
                    subject = parts[0]
                    subtopic = parts[1] if len(parts) >= 3 else "General"
                else:
                    subject = "General"
                    subtopic = "General"

                yield {
                    "full_path": path,
                    "subject": subject,
                    "subtopic": subtopic,
                    "filename": path.name,
                }
            except ValueError as e:
                print(f"[WARN] Could not process path {path}: {e}")
                continue


def extract_text(path: Path) -> str:
    """
    Extract raw text from a supported file.

    Args:
        path: Path object to the file

    Returns:
        Extracted text as string

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()

    try:
        if suffix in {".txt", ".md"}:
            return _extract_text_file(path)
        elif suffix == ".docx":
            return _extract_docx(path)
        elif suffix == ".pdf":
            return _extract_pdf(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as e:
        print(f"[ERROR] Failed to extract text from {path.name}: {e}")
        raise


def _extract_text_file(path: Path) -> str:
    """
    Extract text from plain text or markdown files.

    Args:
        path: Path to text file

    Returns:
        File content as string
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        try:
            return path.read_text(encoding="latin-1")
        except Exception:
            # Last resort with error replacement
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        raise Exception(f"Failed to read text file: {e}")


def _extract_docx(path: Path) -> str:
    """
    Extract text from Microsoft Word documents.

    Args:
        path: Path to .docx file

    Returns:
        Document text as string
    """
    try:
        doc = docx.Document(str(path))
        lines = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text)

        return "\n".join(lines)
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX file: {e}")


def _extract_pdf(path: Path) -> str:
    """
    Extract text from PDF files.

    Args:
        path: Path to .pdf file

    Returns:
        PDF text content as string
    """
    try:
        texts = []

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        texts.append(page_text)
                except Exception as e:
                    print(f"[WARN] Could not extract text from page {page_num} in {path.name}: {e}")
                    continue

        return "\n\n".join(texts) if texts else ""
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF file: {e}")


def get_supported_extensions() -> set:
    """
    Get the set of supported file extensions.

    Returns:
        Set of supported file extensions (with leading dots)
    """
    return SUPPORTED_EXTENSIONS.copy()


def is_supported_file(path: Path) -> bool:
    """
    Check if a file is supported for text extraction.

    Args:
        path: Path to file

    Returns:
        True if file is supported, False otherwise
    """
    return path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS